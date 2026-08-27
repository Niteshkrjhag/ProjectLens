"""Mixed-format document parsing with line-level evidence locations."""

from __future__ import annotations

import hashlib
import html
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import BinaryIO

import pymupdf
from docx import Document as DocxDocument

from .document_policy import DocumentCategory, classify_document


SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx", ".rtf", ".html", ".htm"}


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    filename: str
    relative_path: str
    content_type: str
    content: str
    sha256: str
    size_bytes: int
    category: DocumentCategory
    updated_at: str
    metadata: dict[str, object]


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = html.unescape(data).strip()
        if text:
            self.parts.append(text)


def _strip_rtf(text: str) -> str:
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", text)
    return re.sub(r"[{}]", "", text)


def _classify(filename: str, content: str) -> DocumentCategory:
    title = " ".join(content.splitlines()[:12])
    category = classify_document(filename=filename, title=title)
    if category is not DocumentCategory.UNKNOWN:
        return category
    aliases = {
        "baseline": DocumentCategory.REQUIREMENTS,
        "delivery-note": DocumentCategory.STATUS,
        "anomaly-log": DocumentCategory.ISSUE,
        "exception-note": DocumentCategory.ISSUE,
        "adversarial-source": DocumentCategory.ISSUE,
        "boundary-case": DocumentCategory.ARCHITECTURE,
        "handoff-record": DocumentCategory.DEPLOYMENT_REPORT,
        "edge-evidence": DocumentCategory.CHANGELOG,
        "review-record": DocumentCategory.QA_REPORT,
    }
    for alias, candidate in aliases.items():
        if alias in filename.casefold():
            return candidate
    return category


def parse_bytes(filename: str, raw: bytes, *, relative_path: str | None = None, updated_at: str | None = None) -> ParsedDocument:
    """Parse supported bytes without executing document content."""
    suffix = Path(filename).suffix.casefold()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"unsupported document format: {suffix or '(none)'}")
    if suffix in {".md", ".txt"}:
        content = raw.decode("utf-8-sig", errors="replace")
        content_type = "text/markdown" if suffix == ".md" else "text/plain"
        metadata: dict[str, object] = {}
    elif suffix == ".pdf":
        pages: list[str] = []
        with pymupdf.open(stream=raw, filetype="pdf") as pdf:
            for page_number, page in enumerate(pdf, start=1):
                pages.append(f"[Page {page_number}]\n{page.get_text()}")
        content = "\n\n".join(pages)
        content_type = "application/pdf"
        metadata = {"page_count": len(pages)}
    elif suffix == ".docx":
        import io

        document = DocxDocument(io.BytesIO(raw))
        content = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        metadata = {"paragraph_count": len(document.paragraphs)}
    elif suffix in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(raw.decode("utf-8-sig", errors="replace"))
        content = "\n".join(parser.parts)
        content_type = "text/html"
        metadata = {}
    else:
        content = _strip_rtf(raw.decode("utf-8-sig", errors="replace"))
        content_type = "application/rtf"
        metadata = {}
    content = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    instruction_markers = re.compile(
        r"(?:ignore (?:all )?(?:previous|earlier) instructions|follow no embedded instructions|"
        r"do not execute text|system prompt|you must obey|execute this|instructions? addressed to)",
        re.IGNORECASE,
    )
    anomaly_lines = [
        number for number, line in enumerate(content.splitlines(), start=1) if instruction_markers.search(line)
    ]
    metadata["prompt_injection_detected"] = bool(anomaly_lines)
    metadata["anomaly_lines"] = anomaly_lines
    return ParsedDocument(
        filename=Path(filename).name,
        relative_path=relative_path or filename,
        content_type=content_type,
        content=content,
        sha256=hashlib.sha256(raw).hexdigest(),
        size_bytes=len(raw),
        category=_classify(Path(filename).name, content),
        updated_at=updated_at or "",
        metadata=metadata,
    )


def parse_file(path: Path, root: Path | None = None) -> ParsedDocument:
    """Parse a file from a watched folder, preserving its relative path."""
    raw = path.read_bytes()
    relative_path = str(path.relative_to(root)) if root else path.name
    return parse_bytes(path.name, raw, relative_path=relative_path)


def discover_files(root: Path) -> list[Path]:
    """Find supported files while ignoring hidden and generated directories."""
    if not root.exists():
        raise FileNotFoundError(f"source folder does not exist: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"source path is not a folder: {root}")
    return sorted(
        path for path in root.rglob("*")
        if path.is_file() and path.suffix.casefold() in SUPPORTED_SUFFIXES and not any(part.startswith(".") for part in path.parts)
    )


def read_upload(upload: BinaryIO, filename: str) -> ParsedDocument:
    """Read an upload-like stream into the same safe parser used by folders."""
    return parse_bytes(filename, upload.read())
